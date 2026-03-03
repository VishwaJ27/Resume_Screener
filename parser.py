import os
import re
import pdfplumber
import docx


#  Helper : clean up messy extracted text

def clean_text(text: str) -> str:
    # remove weird unicode chars
    text = text.encode("ascii", "ignore").decode()

    # collapse multiple blank lines into one
    text = re.sub(r"\n{3,}", "\n\n", text)

    # strip leading/trailing whitespace per line
    lines = [line.strip() for line in text.splitlines()]
    text = "\n".join(lines)

    return text.strip()


#  PDF Parser

def extract_from_pdf(filepath: str) -> str:
    text = ""

    try:
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()

                if page_text:
                    text += page_text + "\n"

    except Exception as e:
        print(f"[ERROR] Failed to read PDF: {e}")
        return ""

    return clean_text(text)


#  DOCX Parser

def extract_from_docx(filepath: str) -> str:
    text = ""

    try:
        doc = docx.Document(filepath)

        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"

        # also grab text from tables (some resumes use tables for layout)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + "\n"

    except Exception as e:
        print(f"[ERROR] Failed to read DOCX: {e}")
        return ""

    return clean_text(text)


#  Main Entry Point — auto detect file type

def parse_resume(filepath: str) -> dict:
    if not os.path.exists(filepath):
        return {"error": f"File not found: {filepath}", "text": ""}

    ext = os.path.splitext(filepath)[1].lower()
    filename = os.path.basename(filepath)

    print(f"\n[INFO] Parsing resume: {filename}")

    if ext == ".pdf":
        raw_text = extract_from_pdf(filepath)

    elif ext == ".docx":
        raw_text = extract_from_docx(filepath)

    else:
        return {"error": f"Unsupported file type: {ext}", "text": ""}

    if not raw_text:
        return {"error": "Could not extract any text", "text": ""}

    result = {
        "filename": filename,
        "file_type": ext,
        "text": raw_text,
        "word_count": len(raw_text.split()),
        "char_count": len(raw_text),
        "error": None
    }

    print(f"[INFO] Extracted {result['word_count']} words from {filename}")
    return result


#  Parse Multiple Resumes at once

def parse_multiple_resumes(folder_path: str) -> list:
    results = []

    supported = (".pdf", ".docx")
    files = [f for f in os.listdir(folder_path) if f.endswith(supported)]

    if not files:
        print("[WARN] No PDF or DOCX files found in folder.")
        return []

    print(f"\n[INFO] Found {len(files)} resume(s) to process...")

    for filename in files:
        full_path = os.path.join(folder_path, filename)
        parsed = parse_resume(full_path)
        results.append(parsed)

    return results


#  Quick test — run this file directly

if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print("Usage: python parser.py <path_to_resume.pdf or .docx>")
        print("       python parser.py <folder_with_resumes>")
        sys.exit(1)

    path = sys.argv[1]

    # folder mode
    if os.path.isdir(path):
        resumes = parse_multiple_resumes(path)
        for r in resumes:
            print(f"\n{'='*50}")
            print(f"File     : {r['filename']}")
            print(f"Words    : {r.get('word_count', 'N/A')}")
            print(f"Error    : {r.get('error', 'None')}")
            print(f"Preview  :\n{r['text'][:300]}...")

    # single file mode
    else:
        result = parse_resume(path)
        print(f"\n{'='*50}")
        print(f"File     : {result['filename']}")
        print(f"Words    : {result.get('word_count', 'N/A')}")
        print(f"Error    : {result.get('error', 'None')}")
        print(f"\nFull Text:\n{result['text']}")